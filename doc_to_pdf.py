from docx import Document
import re
import pandas as pd
from pandas.core.frame import DataFrame
import openpyxl


wb = openpyxl.load_workbook(filename='sample_data.xlsx')

# print(wb.worksheets[0])
ws = wb["Sheet1"]
# sheet = wb.active

for cell in ws["Name"]:
    print(cell.value)
# cell_value = sheet.cell(row=1, column=1)
# print(cell_value.value)

# m_row = sheet.max_row
# print(m_row.value)
# m_col = sheet.max_column

# for i in range(1, m_row+1):
#     cell = sheet.cell(row=i, column=1)
#     print(cell.value)


# print(sheet.rows)

# for row in ws.rows:
#     for cell in row:
#         print(cell.value)

# for col in range(1, 10):
#     print(ws.cell(column=col, row=1).value)

# first_cell_value = "1"
# row = 2
# table = dict()
# while first_cell_value != None:
#     if not ws.row_dimensions[row].hidden:
#         table[row] = dict()
#         for col in range(1, 10):
#             table[row][col] = ws.cell(column=col, row=row).value
#     row += 1
#     first_cell_value = ws.cell(column=1, row=row).value
#     print(first_cell_value)
# filename = "sample.docx"


# doc = Document(filename)


# # single_para = document.paragraphs


# for run in doc.paragraphs:
#     print(run.text)

# pin_code = run.text
# # all_data = DataFrame()
# filename = "sample_data.xlsx"
# df = pd.read_excel(filename)
# name = df['Name'].values.tolist()
# pincode = df['Pincode'].values.tolist()

# # print(pincode)
# for i, j in zip(name, pincode):
#     d = {"Name": i,
#          "Pincode": j}

#     # print(d["Pincode"])

#     document = Document()

#     p = document.add_paragraph(pin_code + str(d["Pincode"]))

#     print(p)

# document.save('demo.docx')

# print(d)

# print(df["Name"])
# print
# d = df.to_dict()
# print(df.drop(0, axis=0))

# for index, row in df.iterrows():
# print(index, row["Name"])
# for i in d:
# print(i)
# print(df["Pincode"])
# all_data = all_data.append(d, ignore_index=True)
# print(all_data.Name)

# for p in document.paragraphs[1]:
#     # txt = p.text
#     print(p.text)
# p.text = "Name:"
#     # x = re.search("^antiopam*ipsum$", txt)

#     # TITLE = r"(?:[A-Z][a-z]*\.\s*)?"

#     # r = re.findall(TITLE, txt)
#     # print(r)
#     # print(x)
